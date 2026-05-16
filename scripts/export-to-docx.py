#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
导出为Word文档脚本
将Markdown格式的汇报文档转换为Word格式
"""

import sys
from pathlib import Path

def export_to_docx(markdown_content: str, output_path: str) -> bool:
    """
    将Markdown内容导出为Word文档
    
    Args:
        markdown_content: Markdown格式的内容
        output_path: 输出文件路径
        
    Returns:
        是否成功
    """
    try:
        # 尝试导入python-docx
        from docx import Document
        from docx.shared import Pt, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        print("错误：需要安装python-docx库")
        print("请运行：pip install python-docx")
        return False
    
    # 创建文档
    doc = Document()
    
    # 设置默认字体
    style = doc.styles['Normal']
    style.font.name = '微软雅黑'
    style.font.size = Pt(11)
    
    # 解析Markdown内容
    lines = markdown_content.split('\n')
    i = 0
    while i < len(lines):
        line = lines[i].strip()
        
        # 跳过空行
        if not line:
            i += 1
            continue
        
        # 一级标题
        if line.startswith('# '):
            doc.add_heading(line[2:], level=1)
        
        # 二级标题
        elif line.startswith('## '):
            doc.add_heading(line[3:], level=2)
        
        # 三级标题
        elif line.startswith('### '):
            doc.add_heading(line[4:], level=3)
        
        # 四级标题
        elif line.startswith('#### '):
            doc.add_heading(line[5:], level=4)
        
        # 引用块
        elif line.startswith('> '):
            p = doc.add_paragraph()
            p.style = 'Quote'
            p.add_run(line[2:])
        
        # 无序列表
        elif line.startswith('- '):
            doc.add_paragraph(line[2:], style='List Bullet')
        
        # 有序列表
        elif line and line[0].isdigit() and '. ' in line:
            text = line.split('. ', 1)[1]
            doc.add_paragraph(text, style='List Number')
        
        # 代码块
        elif line.startswith('```'):
            # 收集代码块内容
            code_lines = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith('```'):
                code_lines.append(lines[i])
                i += 1
            code_text = '\n'.join(code_lines)
            p = doc.add_paragraph()
            p.add_run(code_text).font.name = 'Consolas'
        
        # 表格（简化处理）
        elif line.startswith('|'):
            # 收集表格行
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith('|'):
                if '---' not in lines[i]:  # 跳过分隔线
                    table_lines.append(lines[i])
                i += 1
            i -= 1  # 回退一行
            
            # 解析表格
            if table_lines:
                rows = []
                for tl in table_lines:
                    cells = [c.strip() for c in tl.split('|')[1:-1]]
                    rows.append(cells)
                
                if rows:
                    # 创建表格
                    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
                    table.style = 'Table Grid'
                    
                    for row_idx, row_data in enumerate(rows):
                        for col_idx, cell_data in enumerate(row_data):
                            table.rows[row_idx].cells[col_idx].text = cell_data
        
        # 普通段落
        else:
            # 处理加粗
            if '**' in line:
                p = doc.add_paragraph()
                parts = line.split('**')
                for j, part in enumerate(parts):
                    run = p.add_run(part)
                    if j % 2 == 1:
                        run.bold = True
            else:
                doc.add_paragraph(line)
        
        i += 1
    
    # 保存文档
    doc.save(output_path)
    print(f"Word文档已保存到：{output_path}")
    return True


def convert_file(input_path: str, output_path: str = None) -> bool:
    """
    将Markdown文件转换为Word文档
    
    Args:
        input_path: 输入的Markdown文件路径
        output_path: 输出的Word文件路径（可选）
        
    Returns:
        是否成功
    """
    input_path = Path(input_path)
    
    if not input_path.exists():
        print(f"错误：文件不存在 - {input_path}")
        return False
    
    if output_path is None:
        output_path = input_path.with_suffix('.docx')
    
    markdown_content = input_path.read_text(encoding='utf-8')
    return export_to_docx(markdown_content, str(output_path))


# 命令行使用
if __name__ == "__main__":
    if len(sys.argv) < 2:
        print("用法：python export-to-docx.py <markdown文件> [输出文件]")
        print("示例：python export-to-docx.py report.md report.docx")
        sys.exit(1)
    
    input_file = sys.argv[1]
    output_file = sys.argv[2] if len(sys.argv) > 2 else None
    
    convert_file(input_file, output_file)
